#!/usr/bin/env python3
"""
Generates AgencyLead-Radar-Setup-Guide.docx — the beginner-facing, step-by-step
connection guide for a new owner. Re-run after changing env vars or integrations:

    python3 scripts/build_setup_guide.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x1D, 0x4E, 0xD8)
SLATE = RGBColor(0x33, 0x41, 0x55)
GREY = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0xB4, 0x1C, 0x1C)

doc = Document()

# ---------------------------------------------------------------- base styles
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for name, size, colour, before in (
    ('Heading 1', 20, BLUE, 22),
    ('Heading 2', 14, SLATE, 16),
    ('Heading 3', 11.5, SLATE, 12),
):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = colour
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(6)


def shade(cell, hexcolour):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexcolour)
    cell._tc.get_or_add_tcPr().append(el)


def p(text='', bold=False, italic=False, size=None, colour=None, style=None,
      space_after=None, align=None):
    par = doc.add_paragraph(style=style)
    if text:
        run = par.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if colour:
            run.font.color.rgb = colour
    if space_after is not None:
        par.paragraph_format.space_after = Pt(space_after)
    if align:
        par.alignment = align
    return par


def rich(parts, style=None):
    """parts = [(text, {'bold':True,'code':True,...}), ...]"""
    par = doc.add_paragraph(style=style)
    for text, opts in parts:
        run = par.add_run(text)
        run.bold = opts.get('bold', False)
        run.italic = opts.get('italic', False)
        if opts.get('code'):
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x9A, 0x34, 0x12)
        if opts.get('colour'):
            run.font.color.rgb = opts['colour']
    return par


def bullet(text, level=0):
    par = doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')
    par.paragraph_format.space_after = Pt(3)
    return par


def step(n, text):
    par = doc.add_paragraph()
    r = par.add_run(f'{n}.  ')
    r.bold = True
    r.font.color.rgb = BLUE
    par.add_run(text)
    par.paragraph_format.left_indent = Inches(0.25)
    par.paragraph_format.space_after = Pt(4)
    return par


def code(text):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    par.paragraph_format.left_indent = Inches(0.3)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    return par


def callout(title, body, fill='FEF3C7'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    shade(cell, fill)
    cell.text = ''
    par = cell.paragraphs[0]
    r = par.add_run(title + '  ')
    r.bold = True
    par.add_run(body)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def table(headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], 'E2E8F0')
        par = hdr[i].paragraphs[0]
        run = par.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            par = cells[i].paragraphs[0]
            run = par.add_run(str(val))
            run.font.size = Pt(9.5)
            if val.startswith(('sk_', 'whsec_', 'price_', 're_', 'DATABASE', 'APP_', 'NEXT_',
                               'OPENAI', 'STRIPE', 'RESEND', 'EMAIL_', 'ADMIN_', 'SUPERADMIN')):
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
    if widths:
        for r in tbl.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def check(text):
    par = doc.add_paragraph()
    r = par.add_run('☐  ')
    r.font.size = Pt(12)
    par.add_run(text)
    par.paragraph_format.left_indent = Inches(0.25)
    par.paragraph_format.space_after = Pt(4)
    return par


# --------------------------------------------------------------- title page
t = p('AgencyLead Radar', align=WD_ALIGN_PARAGRAPH.CENTER)
t.runs[0].font.size = Pt(34)
t.runs[0].bold = True
t.runs[0].font.color.rgb = BLUE
t.paragraph_format.space_before = Pt(120)

s = p('Complete Setup & Handover Guide', align=WD_ALIGN_PARAGRAPH.CENTER)
s.runs[0].font.size = Pt(17)
s.runs[0].font.color.rgb = SLATE

s2 = p('Written for a non-technical owner. No coding required.',
       align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, colour=GREY, size=11)
s2.paragraph_format.space_after = Pt(40)

p('This guide takes you from "I have been given a GitHub repository" to a fully '
  'live, paying-customer-ready SaaS. Every service you need to connect — the '
  'database, OpenAI, Stripe and Resend — has its own chapter with numbered steps, '
  'exactly what to copy, and how to prove it worked.',
  align=WD_ALIGN_PARAGRAPH.CENTER, colour=SLATE)

doc.add_page_break()

# ------------------------------------------------------------------ contents
doc.add_heading('What is in this guide', level=1)
table(['Part', 'What you will do', 'Time'], [
    ['0. Before you start', 'What the accounts cost and what you need to hand', '5 min'],
    ['1. GitHub', 'Take ownership of the code', '10 min'],
    ['2. Vercel', 'Get the site hosted on the internet', '15 min'],
    ['3. Database', 'Create the Postgres database that stores everything', '10 min'],
    ['4. Secrets & admin logins', 'Lock down the app and set your own passwords', '10 min'],
    ['5. First deploy', 'Put it live and log in for the first time', '10 min'],
    ['6. OpenAI', 'Switch on AI audits and Find Prospects', '15 min'],
    ['7. Stripe', 'Take real money for subscriptions', '45 min'],
    ['8. Resend', 'Send welcome and waitlist emails from your own domain', '30 min'],
    ['9. Custom domain', 'Move off .vercel.app to your own name', '20 min'],
    ['10. Super admin', 'Run the business: customers, plans, MRR, dates', '—'],
    ['11. Go-live checklist', 'The final tick-list before you take a customer', '—'],
    ['12. Reference', 'Every environment variable, costs, troubleshooting', '—'],
    ['13. Honest disclosures', 'What this product is and is not, stated plainly', '—'],
], widths=[1.5, 3.9, 0.8])

doc.add_page_break()

# ------------------------------------------------------- 0. before you start
doc.add_heading('Part 0 — Before you start', level=1)

p('AgencyLead Radar is a Next.js web application. You will never need to open a '
  'code editor or type a command into a terminal. Everything in this guide is done '
  'by clicking around in websites and copying values between them.')

doc.add_heading('The four accounts you will create', level=2)
table(['Service', 'What it does for you', 'Required?', 'Cost to start'], [
    ['Vercel', 'Hosts the website and runs the code', 'Yes', 'Free (Hobby)'],
    ['Neon Postgres', 'The database — users, leads, audits, waitlist', 'Yes', 'Free tier'],
    ['OpenAI', 'AI audits and the Find Prospects feature', 'Strongly recommended', 'Pay per use, ~$0.0002 per audit'],
    ['Stripe', 'Takes card payments for subscriptions', 'Only if you want revenue', '2.9% + 30¢ per charge'],
    ['Resend', 'Sends welcome and waitlist emails', 'Optional', 'Free up to 3,000/month'],
], widths=[1.1, 2.7, 1.1, 1.5])

callout('Important:',
        'OpenAI, Stripe and Resend are all genuinely optional. The app is built so that '
        'if a key is missing, that one feature politely switches itself off and everything '
        'else keeps working. You can go live first and add them one at a time.')

doc.add_heading('What you need to hand', level=2)
bullet('A GitHub account (free) — github.com')
bullet('A credit or debit card — for OpenAI credit and, if you use it, Stripe identity verification. Vercel, Neon and Resend all start free.')
bullet('A domain name, if you already own one. You can skip this and use the free yourname.vercel.app address at first.')
bullet('About two hours if you do everything in one sitting.')

doc.add_page_break()

# ------------------------------------------------------------- 1. github
doc.add_heading('Part 1 — Take ownership of the code on GitHub', level=1)

p('The entire product lives in one GitHub repository. Whoever owns that repository '
  'owns the product. This is the first thing to transfer.')

step(1, 'Create a free account at github.com if you do not have one, and confirm your email address.')
step(2, 'Send the seller your GitHub username and ask them to transfer the repository to you. In GitHub this is Settings → General → scroll to the bottom → "Transfer ownership". A transfer is better than them simply adding you as a collaborator, because it makes you the owner.')
step(3, 'Accept the transfer email. The repository now appears under your own account.')
step(4, 'Open the repository and check you can see folders named app, lib and prisma, and a file named package.json. If you can, you have the whole product.')

callout('Do not skip this.',
        'If the repository stays in the seller\'s account they can delete or change your '
        'product at any time, and you cannot deploy updates. Get the transfer done before '
        'you pay for anything else.', fill='FEE2E2')

doc.add_heading('What the folders are, in plain English', level=2)
table(['Folder', 'What lives there'], [
    ['app', 'Every page you see in the browser and every API endpoint behind them'],
    ['lib', 'Shared logic — login, plans and quotas, Stripe, email, rate limiting'],
    ['prisma', 'The database structure, plus the seed script that creates your admin accounts'],
    ['components', 'Reusable buttons, cards and form pieces'],
    ['.env.example', 'A commented list of every setting the app accepts — your cheat sheet'],
], widths=[1.3, 5.0])

doc.add_page_break()

# -------------------------------------------------------------- 2. vercel
doc.add_heading('Part 2 — Put the site on the internet with Vercel', level=1)

p('Vercel is the hosting company. It watches your GitHub repository, and every time '
  'the code changes it automatically rebuilds and republishes the site. It is made by '
  'the same team that makes Next.js, which is what this app is written in.')

step(1, 'Go to vercel.com and click "Sign Up". Choose "Continue with GitHub" — this links the two accounts for you.')
step(2, 'On your Vercel dashboard click "Add New…" → "Project".')
step(3, 'Find your agencylead repository in the list and click "Import". If you do not see it, click "Adjust GitHub App Permissions" and grant Vercel access to it.')
step(4, 'On the configuration screen, leave every setting exactly as it is. Framework Preset should already say "Next.js". Do not change the build command.')
step(5, 'Do NOT click Deploy yet. Scroll down to "Environment Variables" — we need a database first. Click "Deploy" only after Part 4.')

callout('If you already clicked Deploy and it failed:',
        'That is completely normal and harmless. The build fails because there is no '
        'database yet. Carry on with Part 3, then redeploy at the end of Part 4.')

doc.add_heading('The one Vercel screen you will keep coming back to', level=2)
p('Almost everything in this guide involves adding a value to the same place:')
code('Vercel → your project → Settings → Environment Variables')
p('Each entry has a name (like OPENAI_API_KEY) and a value (the secret you copied). '
  'Tick all three environments — Production, Preview and Development — unless this guide '
  'says otherwise. Remember the golden rule below.')

callout('The golden rule of environment variables:',
        'Changing an environment variable does NOTHING until you redeploy. After you add or '
        'edit one, always go to the Deployments tab, click the "…" menu on the newest '
        'deployment and choose "Redeploy". This catches out almost everybody once.',
        fill='DBEAFE')

doc.add_page_break()

# ------------------------------------------------------------ 3. database
doc.add_heading('Part 3 — Create the database', level=1)

p('The database is where your customers, their leads, the AI audits and the waitlist '
  'signups are permanently stored. This app uses Postgres. The easiest option is Neon, '
  'which you can add from inside Vercel in about two minutes.')

step(1, 'In your Vercel project, open the "Storage" tab.')
step(2, 'Click "Create Database", then choose "Neon — Serverless Postgres" from the marketplace list.')
step(3, 'Accept the free plan, pick the region closest to your customers (for US agencies choose a US East region), and click Create.')
step(4, 'When Vercel asks whether to connect the database to this project, say yes and connect it to all environments.')
step(5, 'Go to Settings → Environment Variables and confirm a variable called DATABASE_URL now exists. Vercel adds it for you.')

callout('If DATABASE_URL is missing or you used a different provider:',
        'Open your database provider\'s dashboard, copy the "connection string" — it starts '
        'with postgresql:// — and add it to Vercel manually with the name DATABASE_URL. '
        'Make sure the string ends with ?sslmode=require.')

doc.add_heading('What happens to the database automatically', level=2)
p('You never have to create tables by hand. Every time Vercel builds the site it runs '
  'this build command for you:')
code('prisma generate && prisma db push && prisma db seed && next build')
bullet('prisma db push creates or updates all the tables to match the app.')
bullet('prisma db seed creates your admin and super admin accounts and loads the demo leads used by the public /demo page.')

callout('Your real customer data is safe.',
        'The seed script only creates records that are missing and re-applies your admin '
        'passwords. It never deletes customers, leads or payments on redeploy.', fill='DCFCE7')

doc.add_page_break()

# ------------------------------------------------ 4. secrets & admin logins
doc.add_heading('Part 4 — Secrets and your admin logins', level=1)

doc.add_heading('4.1  APP_SECRET — the key that keeps logins secure', level=2)
p('When somebody logs in, the app gives their browser a signed pass. APP_SECRET is the '
  'signature key. If it is missing the app refuses to start in production, on purpose — '
  'a guessable session key would let anyone impersonate your customers.')

step(1, 'Generate a long random string. The simplest way with no software: visit a reputable random-string generator, or use 1Password/Bitwarden\'s password generator set to 40+ characters.')
step(2, 'If you are comfortable with a terminal on your own computer, this is the ideal command:')
code('openssl rand -base64 32')
step(3, 'In Vercel → Settings → Environment Variables, add:')
table(['Name', 'Value'], [['APP_SECRET', 'the long random string you just generated']], widths=[2.0, 4.3])
step(4, 'Save it in your password manager as well. If you ever change it, every logged-in customer is signed out — annoying but not dangerous.')

callout('Never share APP_SECRET.',
        'It must not appear in emails, screenshots, support chats or the GitHub repository. '
        'It lives only in Vercel and your password manager.', fill='FEE2E2')

doc.add_heading('4.2  Your admin and super admin accounts', level=2)
p('The app ships with two staff accounts. Their passwords have publicly known defaults, '
  'so you must override them before you go live.')

table(['Variable', 'What it sets', 'Default (change it!)'], [
    ['ADMIN_EMAIL', 'Login for the admin dashboard', 'admin@agencyleadradar.com'],
    ['ADMIN_PASSWORD', 'Password for the admin account', 'admin123'],
    ['SUPERADMIN_EMAIL', 'Login for the whole-site super admin page', 'superadmin@agencyleadradar.com'],
    ['SUPERADMIN_PASSWORD', 'Password for the super admin account', 'superadmin123'],
], widths=[1.7, 2.7, 1.9])

step(1, 'Add all four variables in Vercel with your own email addresses and strong, unique passwords (16+ characters).')
step(2, 'Store both logins in your password manager.')

callout('How password rotation works here:',
        'The seed script re-applies ADMIN_PASSWORD and SUPERADMIN_PASSWORD on every single '
        'deploy. So to change a staff password later, just change the variable in Vercel and '
        'redeploy — do not change it inside the app, or the next deploy will overwrite it.',
        fill='DBEAFE')

doc.add_heading('4.3  Your public address', level=2)
p('One more variable tells the app what its own web address is. It is used in email links '
  'and Stripe redirects.')
table(['Name', 'Value'], [
    ['NEXT_PUBLIC_APP_URL', 'https://your-project.vercel.app  (include the https://)'],
], widths=[2.0, 4.3])
p('Come back and change this to your custom domain in Part 9.')

doc.add_page_break()

# ---------------------------------------------------------- 5. first deploy
doc.add_heading('Part 5 — Your first deploy', level=1)

p('You now have the four things the app truly requires: DATABASE_URL, APP_SECRET, your '
  'admin credentials, and NEXT_PUBLIC_APP_URL. Time to go live.')

step(1, 'In Vercel, open the "Deployments" tab.')
step(2, 'On the newest deployment, click the "…" menu → "Redeploy" → confirm.')
step(3, 'Watch the build log. It takes two to four minutes. You are looking for "Compiled successfully" followed by "Deployment completed".')
step(4, 'Click "Visit". The marketing homepage should load.')

doc.add_heading('Prove it actually works', level=2)
table(['Go to', 'You should see'], [
    ['/', 'The homepage with pricing and a call to action'],
    ['/demo', 'A working demo dashboard preloaded with sample leads'],
    ['/signup', 'A signup form — create a test account here'],
    ['/dashboard', 'Your own empty dashboard after signing up'],
    ['/login', 'Log in with your ADMIN_EMAIL and password'],
    ['/admin', 'The admin dashboard with waitlist signups'],
    ['/admin/overview', 'The super admin page (log in as SUPERADMIN_EMAIL for this one)'],
], widths=[1.5, 4.8])

callout('If you get a 404 on every page:',
        'Check that Vercel is deploying the "main" branch. Settings → Git → Production Branch '
        'must say main. This is the single most common cause of a completely blank site.',
        fill='FEE2E2')

doc.add_page_break()

# ------------------------------------------------------------- 6. openai
doc.add_heading('Part 6 — Connect OpenAI (AI audits and Find Prospects)', level=1)

p('OpenAI powers the two features customers actually pay for:')
bullet('AI Audit — writes a plain-English website critique and a ready-to-send outreach email for a lead.')
bullet('Find Prospects — asks the user a few questions about who they want to work with, then suggests businesses to research.')
p('Without an OpenAI key, both buttons explain that AI is not configured and everything '
  'else in the app carries on working.')

doc.add_heading('6.1  Create the key', level=2)
step(1, 'Go to platform.openai.com and sign up. Note this is the developer platform — it is a separate thing from a ChatGPT Plus subscription, and ChatGPT Plus does not include API access.')
step(2, 'Open Settings → Billing and add a payment method. Add $10 of credit to begin with; that is a very large amount of usage for this app.')
step(3, 'While in Billing, set a monthly usage limit — $20 is a sensible starting cap. This is your protection against a runaway bill.')
step(4, 'Go to API keys → "Create new secret key". Name it "AgencyLead Radar Production".')
step(5, 'Copy the key immediately. It starts with sk- and OpenAI will never show it to you again.')

doc.add_heading('6.2  Add it to Vercel', level=2)
table(['Name', 'Value'], [['OPENAI_API_KEY', 'sk-…  (the key you just copied)']], widths=[2.0, 4.3])
step(1, 'Add the variable, tick all three environments, and save.')
step(2, 'Redeploy (Deployments → "…" → Redeploy). Remember the golden rule.')

doc.add_heading('6.3  Test it', level=2)
step(1, 'Log in as your test customer and go to the Lead Scanner.')
step(2, 'Add a lead — any real local business name and website will do.')
step(3, 'Open the lead and click "Generate AI Audit". Within about ten seconds you should get a score breakdown and a draft outreach email.')
step(4, 'Back on the leads page, click "Find Prospects", answer the questions, and check you get a list of suggested businesses.')

doc.add_heading('6.4  What it costs, honestly', level=2)
p('The app uses gpt-4o-mini, the cheap fast model. A single audit costs roughly two '
  'hundredths of a cent. A thousand audits a month is well under $1. Your OpenAI bill '
  'will be a rounding error next to your subscription revenue.')

callout('About the AI\'s suggestions:',
        'The model cannot browse the web. Find Prospects therefore returns businesses to '
        'research, not verified facts — and the app marks every imported suggestion as '
        '"unverified" with an amber banner until the user confirms it. Do not remove that '
        'warning; it is what keeps the feature honest.')

doc.add_page_break()

# ------------------------------------------------------------- 7. stripe
doc.add_heading('Part 7 — Connect Stripe (take real money)', level=1)

p('This is the longest chapter because it is the one that pays you. Take it slowly and '
  'do the whole thing in Test mode first. There is a switch at the top of the Stripe '
  'dashboard labelled "Test mode" — leave it ON until the very end.')

doc.add_heading('7.1  What the app expects', level=2)
p('There are four paid tiers in the app. Free needs no Stripe product; the other three do:')
table(['Plan', 'Price', 'Leads / month', 'AI audits / month', 'Vercel variable'], [
    ['Free', '$0', '10', '3', '— none —'],
    ['Starter', '$29', '100', '50', 'STRIPE_PRICE_STARTER'],
    ['Agency', '$79', '500', '250', 'STRIPE_PRICE_AGENCY'],
    ['Pro', '$149', '2,000', '1,000', 'STRIPE_PRICE_PRO'],
], widths=[0.85, 0.7, 1.05, 1.15, 2.55])

doc.add_heading('7.2  Create your Stripe account', level=2)
step(1, 'Sign up at stripe.com. Choose the country your business is registered in — this cannot be changed later.')
step(2, 'Complete the business verification (legal name, address, bank account, ID). Stripe will not release payouts until this is done, though you can build and test everything first.')
step(3, 'Confirm "Test mode" is switched ON in the top-right of the dashboard.')

doc.add_heading('7.3  Create the three products', level=2)
p('Repeat these steps three times — once for Starter, once for Agency, once for Pro.')
step(1, 'Go to Product catalogue → "Add product".')
step(2, 'Name it exactly as the plan is named in the app: Starter, Agency or Pro.')
step(3, 'Under Pricing choose: Recurring, Monthly, USD, and the amount from the table above ($29, $79 or $149).')
step(4, 'Save the product.')
step(5, 'On the product page find the Pricing section and copy the Price ID. It looks like price_1QxAbCDeFgHiJkLm. Copy the PRICE id, not the product id — a product id starts with prod_ and will not work.')

callout('The single most common Stripe mistake:',
        'Pasting a prod_… id where a price_… id is expected. Checkout will fail every time. '
        'If checkout returns an error, check this first.', fill='FEE2E2')

doc.add_heading('7.4  Get your secret key', level=2)
step(1, 'Go to Developers → API keys.')
step(2, 'Reveal and copy the "Secret key". In Test mode it starts with sk_test_. In Live mode it starts with sk_live_.')

doc.add_heading('7.5  Create the webhook', level=2)
p('A webhook is how Stripe tells your app "this person just paid" or "this person '
  'cancelled". Without it, customers would pay and never be upgraded. It is not optional.')

step(1, 'Go to Developers → Webhooks → "Add endpoint".')
step(2, 'For the endpoint URL enter your site address followed by the webhook path:')
code('https://your-domain.com/api/stripe/webhook')
step(3, 'Click "Select events" and tick exactly these four:')
table(['Event', 'What it does in the app'], [
    ['checkout.session.completed', 'A customer finished paying — upgrade their account'],
    ['customer.subscription.created', 'Record the new subscription and its plan'],
    ['customer.subscription.updated', 'Plan change, renewal date, or scheduled cancellation'],
    ['customer.subscription.deleted', 'Subscription ended — drop the account back to Free'],
], widths=[2.3, 4.0])
step(4, 'Create the endpoint, then on its page click "Reveal" under Signing secret and copy it. It starts with whsec_.')

callout('Why the signing secret matters:',
        'The app verifies every webhook against this secret and rejects anything that fails. '
        'That is what stops a stranger from sending a fake "payment succeeded" message and '
        'giving themselves a free Pro account. If the secret is wrong, all upgrades silently '
        'stop working.', fill='DBEAFE')

doc.add_heading('7.6  Add all five values to Vercel', level=2)
table(['Name', 'Where it came from'], [
    ['STRIPE_SECRET_KEY', 'Developers → API keys (sk_test_… for now)'],
    ['STRIPE_WEBHOOK_SECRET', 'The webhook endpoint page (whsec_…)'],
    ['STRIPE_PRICE_STARTER', 'The Starter product\'s Price ID (price_…)'],
    ['STRIPE_PRICE_AGENCY', 'The Agency product\'s Price ID (price_…)'],
    ['STRIPE_PRICE_PRO', 'The Pro product\'s Price ID (price_…)'],
], widths=[2.2, 4.1])
p('Add all five, then redeploy.')

doc.add_heading('7.7  Test the full payment flow', level=2)
step(1, 'Sign up as a brand-new test customer on your live site.')
step(2, 'Go to Pricing and choose Starter. You should be sent to a Stripe checkout page.')
step(3, 'Pay with Stripe\'s test card:')
code('Card    4242 4242 4242 4242\nExpiry  any future date, e.g. 12/34\nCVC     any 3 digits\nZIP     any 5 digits')
step(4, 'You should be returned to the app. Check the billing page now says Starter.')
step(5, 'In Stripe, open Developers → Webhooks → your endpoint and confirm the events show a 200 response. A red failure here means the upgrade did not reach your app.')
step(6, 'Log in as super admin and confirm the account shows plan Starter, status active, and a renewal date.')
step(7, 'Back as the customer, click "Manage billing". This opens Stripe\'s own billing portal where they can change card or cancel. Cancel the test subscription and confirm the account returns to Free.')

callout('Only when every step above passes:',
        'Switch Stripe out of Test mode, recreate the three products and the webhook in Live '
        'mode (live mode is a completely separate world — nothing carries over), then replace '
        'STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET and the three price IDs in Vercel with the '
        'live versions and redeploy. Then make one real USD purchase on your own card and '
        'refund it.', fill='DCFCE7')

doc.add_page_break()

# ------------------------------------------------------------- 8. resend
doc.add_heading('Part 8 — Connect Resend (emails)', level=1)

p('Resend sends the app\'s transactional emails. Three are wired up today:')
table(['Email', 'When it sends', 'Goes to'], [
    ['Welcome', 'Someone creates an account', 'The new customer'],
    ['Waitlist confirmation', 'Someone joins the public waitlist', 'The person who signed up'],
    ['New waitlist signup alert', 'Someone joins the waitlist', 'You (optional)'],
], widths=[1.8, 2.4, 2.1])

p('If you skip this chapter the app runs perfectly and simply sends nothing. Sending is '
  'also "fire and forget" — if Resend has an outage, signups still succeed.')

doc.add_heading('8.1  Create your Resend account', level=2)
step(1, 'Go to resend.com and sign up. The free tier covers 3,000 emails a month and 100 a day, which is plenty to start.')
step(2, 'Verify your email address.')

doc.add_heading('8.2  Verify your sending domain', level=2)
p('You can send test emails immediately from Resend\'s shared onboarding address, but for '
  'real customers you must verify your own domain or your emails will land in spam.')

step(1, 'In Resend go to Domains → "Add Domain" and enter your domain, e.g. yourdomain.com.')
step(2, 'Resend shows you a set of DNS records — typically an MX record, a TXT record for SPF, and a TXT record for DKIM.')
step(3, 'Log in to wherever your domain is registered (GoDaddy, Namecheap, Cloudflare, Vercel Domains…) and find the DNS settings.')
step(4, 'Add each record exactly as Resend shows it. Copy and paste — do not retype. If your registrar automatically appends your domain to the name field, enter only the prefix part.')
step(5, 'Go back to Resend and click "Verify". DNS can take anything from five minutes to a few hours to propagate. Keep clicking verify.')
step(6, 'Wait until the domain shows a green "Verified" status before going further.')

callout('What these records actually do:',
        'SPF and DKIM are how Gmail and Outlook check that an email claiming to come from '
        'your domain really did. Without them your welcome emails go straight to spam, and '
        'customers who never see a welcome email are far more likely to churn.')

doc.add_heading('8.3  Create the API key', level=2)
step(1, 'In Resend go to API Keys → "Create API Key".')
step(2, 'Name it "AgencyLead Radar Production". Permission "Sending access" is enough.')
step(3, 'Copy the key. It starts with re_ and is shown only once.')

doc.add_heading('8.4  Add the three variables to Vercel', level=2)
table(['Name', 'Example value', 'Required?'], [
    ['RESEND_API_KEY', 're_AbC123…', 'Yes, to send anything'],
    ['EMAIL_FROM', 'AgencyLead Radar <hello@yourdomain.com>', 'Yes, to send anything'],
    ['ADMIN_NOTIFY_EMAIL', 'you@yourdomain.com', 'Optional'],
], widths=[1.7, 3.2, 1.4])

callout('EMAIL_FROM must use your verified domain.',
        'The format "Display Name <address@domain.com>" is what customers see in their inbox. '
        'If the domain part does not match a verified domain in Resend, every send is rejected '
        'and you will see the failure in your Vercel logs.', fill='DBEAFE')

p('Add all three, then redeploy.')

doc.add_heading('8.5  Test it', level=2)
step(1, 'Sign up on your live site with a personal email address you can check.')
step(2, 'A "Welcome to AgencyLead Radar" email should arrive within about a minute. Check spam if not.')
step(3, 'Submit the public waitlist form with another address. Confirm both the confirmation email arrives and, if you set ADMIN_NOTIFY_EMAIL, that you get the alert.')
step(4, 'In Resend, open the Logs tab. Every attempt appears there with its result — this is where you diagnose anything that did not arrive.')

doc.add_heading('8.6  Emails that are not built yet', level=2)
p('For a future owner\'s roadmap and for honesty with your own customers, the following '
  'are NOT implemented today and would need building:')
bullet('Password reset by email')
bullet('Email address verification at signup')
bullet('Payment receipts from your app (Stripe does send its own receipts — enable them in Stripe → Settings → Emails)')
bullet('Quota warning emails when a customer approaches their monthly limit')

doc.add_page_break()

# ------------------------------------------------------------- 9. domain
doc.add_heading('Part 9 — Point your own domain at the site', level=1)

step(1, 'In Vercel go to your project → Settings → Domains.')
step(2, 'Type your domain, e.g. agencyleadradar.com, and click Add.')
step(3, 'Vercel shows the DNS records to add. Usually an A record for the root domain and a CNAME for www.')
step(4, 'Add those records at your registrar, exactly as shown.')
step(5, 'Wait for Vercel to show "Valid Configuration". HTTPS is issued automatically and free.')

doc.add_heading('Then update these three things', level=2)
check('Change NEXT_PUBLIC_APP_URL in Vercel to https://yourdomain.com and redeploy. If you skip this, email links and Stripe redirects will still point at the old .vercel.app address.')
check('Update the Stripe webhook endpoint URL to https://yourdomain.com/api/stripe/webhook.')
check('Check that EMAIL_FROM uses this same domain.')

doc.add_page_break()

# --------------------------------------------------------- 10. super admin
doc.add_heading('Part 10 — Running the business from the super admin page', level=1)

p('Log in with your SUPERADMIN_EMAIL and go to /admin/overview. This page is where you '
  'operate the whole business.')

doc.add_heading('What you can see', level=2)
bullet('Every account, searchable by name, email or company')
bullet('Filters by plan, role and subscription status')
bullet('Live MRR, total accounts, and recent signups')
bullet('How many leads and AI audits each account owns')

doc.add_heading('What you can edit', level=2)
p('Click any account to open the editor. Every field is editable:')
table(['Field', 'Why you would change it'], [
    ['Name, company, email', 'Fixing a customer typo or a change of address'],
    ['Role', 'Promote a teammate to admin or super admin'],
    ['Plan', 'Manually grant a plan — a comped account or a partner deal'],
    ['Subscription status', 'Reflect an off-Stripe arrangement'],
    ['Renews / ends date', 'Extend a trial or honour a promise'],
    ['Cancels at period end', 'Mark an account as winding down'],
    ['Joined date', 'Correct a migrated record'],
    ['Password', 'Reset a locked-out customer directly — there is no self-serve reset yet'],
], widths=[1.8, 4.5])

callout('Read this before editing a paying customer:',
        'Stripe is the source of truth for anyone who pays by card. If you manually change a '
        'plan or status here, the next Stripe webhook will overwrite your change. Use manual '
        'edits for comped, partner and offline accounts — for real card customers, make the '
        'change in Stripe and let it flow through.', fill='FEE2E2')

doc.add_heading('Safety rails built into the page', level=2)
bullet('You cannot remove your own super admin access.')
bullet('You cannot demote or delete the last remaining super admin — the app will not let you lock yourself out.')
bullet('You cannot delete your own account.')
bullet('Deleting an account removes its leads and audits, but keeps the activity log entries with the user detached, so your audit trail survives.')

doc.add_heading('About the MRR figure', level=2)
p('MRR is calculated live from the accounts on screen — it is the sum of the plan prices '
  'of every account whose subscription is active, trialing or past due. It is not stored '
  'anywhere and not collected from Stripe. Treat it as an at-a-glance indicator; treat '
  'Stripe\'s own dashboard as your accounting truth.')

doc.add_page_break()

# ----------------------------------------------------- 11. go-live checklist
doc.add_heading('Part 11 — Go-live checklist', level=1)

doc.add_heading('Must do before your first real customer', level=2)
for item in [
    'GitHub repository transferred into your own account',
    'APP_SECRET set to your own long random value, saved in a password manager',
    'ADMIN_PASSWORD and SUPERADMIN_PASSWORD changed from their published defaults',
    'Logged in successfully as both admin and super admin',
    'DATABASE_URL connected and the site loads with no errors',
    'NEXT_PUBLIC_APP_URL matches the address customers actually use',
    'Signed up as a test customer end to end and reached the dashboard',
    'Privacy policy and Terms reviewed — they ship as templates with [BRACKETED] placeholders that must be completed and checked by a lawyer',
    'Your real company name, contact address and support email filled into those pages',
]:
    check(item)

doc.add_heading('Before you charge anybody', level=2)
for item in [
    'Stripe business verification completed and bank account connected',
    'Three live-mode products created with correct USD prices',
    'Live webhook endpoint created with all four events, returning 200',
    'All five Stripe variables switched to live values in Vercel, and redeployed',
    'One real purchase made on your own card, confirmed in the super admin page, then refunded',
    'Stripe → Settings → Emails: customer receipts switched on',
]:
    check(item)

doc.add_heading('Recommended within the first week', level=2)
for item in [
    'OpenAI monthly spend limit set',
    'Resend domain verified and a welcome email received in a real inbox',
    'Custom domain live with HTTPS',
    'A written support email address published on the site',
    'A plan for backups — check what your database provider retains on your tier',
]:
    check(item)

doc.add_page_break()

# ------------------------------------------------------------ 12. reference
doc.add_heading('Part 12 — Reference', level=1)

doc.add_heading('12.1  Every environment variable', level=2)
table(['Variable', 'Required?', 'What happens without it'], [
    ['DATABASE_URL', 'Required', 'The app cannot start at all'],
    ['APP_SECRET', 'Required', 'The app refuses to run in production, deliberately'],
    ['NEXT_PUBLIC_APP_URL', 'Required', 'Email links and Stripe redirects point to the wrong place'],
    ['ADMIN_EMAIL', 'Recommended', 'Falls back to the published default address'],
    ['ADMIN_PASSWORD', 'Recommended', 'Falls back to the published default "admin123"'],
    ['SUPERADMIN_EMAIL', 'Recommended', 'Falls back to the published default address'],
    ['SUPERADMIN_PASSWORD', 'Recommended', 'Falls back to the published default "superadmin123"'],
    ['OPENAI_API_KEY', 'Optional', 'AI audits and Find Prospects explain they are unavailable'],
    ['STRIPE_SECRET_KEY', 'Optional', 'Pricing falls back to waitlist links; no checkout'],
    ['STRIPE_WEBHOOK_SECRET', 'With Stripe', 'Customers pay but are never upgraded'],
    ['STRIPE_PRICE_STARTER', 'With Stripe', 'The Starter plan cannot be bought'],
    ['STRIPE_PRICE_AGENCY', 'With Stripe', 'The Agency plan cannot be bought'],
    ['STRIPE_PRICE_PRO', 'With Stripe', 'The Pro plan cannot be bought'],
    ['RESEND_API_KEY', 'Optional', 'No emails are sent; everything else works'],
    ['EMAIL_FROM', 'With Resend', 'No emails are sent'],
    ['ADMIN_NOTIFY_EMAIL', 'Optional', 'You get no alert on waitlist signups'],
], widths=[2.0, 1.1, 3.2])

doc.add_heading('12.2  Roughly what it costs to run', level=2)
table(['Service', 'At zero customers', 'At around 50 customers'], [
    ['Vercel', '$0 (Hobby)', '$20/mo (Pro — needed for commercial use)'],
    ['Neon Postgres', '$0', '$0–19/mo'],
    ['OpenAI', '$0', 'Typically under $5/mo'],
    ['Stripe', '$0', '2.9% + 30¢ per charge'],
    ['Resend', '$0', '$0 (under 3,000 emails)'],
    ['Domain', '~$12/yr', '~$12/yr'],
], widths=[1.3, 1.9, 3.1])

callout('Vercel licensing:',
        'Vercel\'s Hobby plan is for non-commercial use. Once you are charging customers you '
        'need the Pro plan at $20/month. Budget for it from day one.')

doc.add_heading('12.3  Troubleshooting', level=2)
table(['Symptom', 'Most likely cause and fix'], [
    ['Every page is a 404', 'Vercel is deploying the wrong branch. Settings → Git → Production Branch must be main.'],
    ['Build fails mentioning DATABASE_URL', 'The database is not connected. Redo Part 3.'],
    ['Build fails mentioning APP_SECRET', 'APP_SECRET is missing in Vercel. Add it and redeploy.'],
    ['I changed a variable and nothing changed', 'You did not redeploy. Deployments → "…" → Redeploy.'],
    ['AI audit says AI is not configured', 'OPENAI_API_KEY is missing or the deploy predates it. Add it and redeploy.'],
    ['AI audit errors after a while', 'Your OpenAI credit ran out or you hit your spend limit. Top up at platform.openai.com.'],
    ['Checkout button does nothing', 'STRIPE_SECRET_KEY missing, or a price ID is wrong. Check you used price_… not prod_….'],
    ['Customer paid but was not upgraded', 'The webhook is failing. Stripe → Developers → Webhooks and read the error. Usually a wrong URL or a stale STRIPE_WEBHOOK_SECRET.'],
    ['Emails never arrive', 'Check Resend → Logs. Usually the domain is unverified or EMAIL_FROM uses a domain Resend does not know.'],
    ['Emails go to spam', 'Domain not verified, or SPF/DKIM records entered incorrectly. Redo Part 8.2.'],
    ['Locked out of admin', 'Change ADMIN_PASSWORD in Vercel and redeploy — the seed re-applies it every build.'],
], widths=[1.9, 4.4])

doc.add_heading('12.4  Where to get help', level=2)
table(['Service', 'Support'], [
    ['Vercel', 'vercel.com/help — deploys, domains, environment variables'],
    ['Neon', 'neon.tech/docs — database size, backups, connection strings'],
    ['OpenAI', 'help.openai.com — billing, keys, rate limits'],
    ['Stripe', 'support.stripe.com — live chat, genuinely good'],
    ['Resend', 'resend.com/docs — deliverability and DNS'],
], widths=[1.3, 5.0])

doc.add_page_break()

# ----------------------------------------------------- 13. honest disclosures
doc.add_heading('Part 13 — Honest disclosures', level=1)

p('These are stated plainly so there are no surprises after handover. None of them are '
  'defects; they are the accurate shape of what you are taking on.')

doc.add_heading('Commercial', level=2)
bullet('This is a pre-revenue product. It has no customers, no revenue history and no traffic history. Any MRR shown in the app is calculated from whatever accounts exist in your database.')
bullet('The leads on the public /demo page are fictional sample data created to demonstrate the interface. They are not real businesses and not a customer list.')
bullet('No testimonials, case studies or usage statistics ship with the product, because none would be true yet.')

doc.add_heading('Legal', level=2)
bullet('The Privacy Policy and Terms of Service are templates containing [BRACKETED] placeholders. They must be completed with your real company details and reviewed by a qualified lawyer in your jurisdiction before you take a paying customer.')
bullet('You will be storing personal data about your customers and their prospects. Depending on where you and they are, GDPR, CCPA or similar rules apply to you as the operator.')

doc.add_heading('Technical', level=2)
bullet('Rate limiting is held in memory per server instance. It effectively slows down abuse but is not a hard guarantee across a scaled deployment. Moving it to a shared store such as Redis is the natural upgrade if you grow.')
bullet('There is no self-serve password reset. Until one is built, you reset customer passwords yourself from the super admin page.')
bullet('There is no email verification at signup, so an account can be created with an address the person does not own.')
bullet('The AI cannot browse the web. Find Prospects returns businesses to research, clearly labelled unverified, not verified facts about them.')
bullet('MRR is derived on the fly, not stored. It is an indicator, not an accounting record — Stripe is your accounting record.')
bullet('Database schema changes are applied with prisma db push on every deploy. Take a backup before making structural changes to the schema.')

doc.add_heading('Roadmap — the obvious next things to build', level=2)
bullet('Password reset and email verification')
bullet('Real website analysis via the PageSpeed Insights API (a PAGESPEED_API_KEY placeholder already exists)')
bullet('Team accounts and seats')
bullet('Quota warning emails and upgrade prompts at 80% usage')
bullet('CRM integrations — HubSpot, Pipedrive, Zapier')
bullet('A shared rate-limit store for multi-instance scale')

doc.add_paragraph()
closing = p('End of guide. Keep this document with your password manager entries — together '
            'they are everything needed to operate AgencyLead Radar.', italic=True, colour=GREY)
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   'AgencyLead-Radar-Setup-Guide.docx')
doc.save(out)
print('Wrote', out)
